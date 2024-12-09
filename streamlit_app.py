import streamlit as st
from openai import OpenAI
import pandas as pd
from docx import Document
from io import BytesIO
import matplotlib.pyplot as plt

# Initialize session state variables safely
st.session_state.setdefault("generated_memo", "")
st.session_state.setdefault("uploaded_files", None)
st.session_state.setdefault("user_changes", "")
st.session_state.setdefault("additional_content", "")

# Function to refresh the page by clearing the session state
def refresh_page():
    for key in list(st.session_state.keys()):
        del st.session_state[key]

# Access the OpenAI API key securely from Streamlit secrets
client = OpenAI(api_key=st.secrets["openai"]["api_key"])

# Add custom CSS for UI
st.markdown(
    """
    <style>
        body {
            background-color: #E6ECF5;
        }
        .title {
            color: #4E81BD;
            font-size: 32px;
            font-weight: bold;
            text-align: center;
            margin: 10px 0 30px;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# Use Streamlit columns to center the logo
col1, col2, col3 = st.columns([2, 2, 1])
with col2:
    st.image("TalosLogo.png", width=150)

# Display title and description
st.markdown("<div class='title'>Commercial Lending Co-Pilot</div>", unsafe_allow_html=True)

# File uploader
uploaded_files = st.file_uploader(
    "Please upload relevant documents, including personal financial statements, LP memos, and any client communication regarding requested changes.",
    accept_multiple_files=True,
    type=("txt", "md", "pdf", "xlsx", "docx", "csv"),
    key="file_upload",
)

# Text area for user input
user_changes = st.text_area(
    "Please insert any additional instructions.",
    value=st.session_state.get("user_changes", ""),
    placeholder="E.g., Who should the memo be signed from? Are there any particularly important details to highlight in the memo?",
    key="user_changes",
)

def normalize_and_match_columns(df):
    """
    Normalize column names and map them to the expected names based on the provided column titles.
    """
    # Strip whitespace and convert to lowercase
    df.columns = df.columns.str.strip().str.lower()
    column_mapping = {
        "year": "year",
        "debt service coverage ratio": "debt service coverage ratio",
        "minimum dscr covenant": "minimum debt service coverage ratio",
    }
    
    # Attempt to map columns to expected names
    normalized_columns = {col: column_mapping[col] for col in df.columns if col in column_mapping}
    
    # Ensure all required columns are present
    required_columns = {"year", "debt service coverage ratio", "minimum debt service coverage ratio"}
    if not required_columns.issubset(set(normalized_columns.values())):
        raise ValueError(
            "The file must contain 'Year', 'Debt Service Coverage Ratio', and 'Minimum DSCR Covenant' columns."
        )
    
    # Rename columns in DataFrame
    df = df.rename(columns=normalized_columns)
    return df

def generate_dscr_chart(df):
    """
    Generate a DSCR chart based on the DataFrame.
    """
    try:
        # Normalize and match columns
        df = normalize_and_match_columns(df)
        
        # Plot the data
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.bar(
            df["year"], 
            df["debt service coverage ratio"], 
            label="Debt Service Coverage Ratio", 
            alpha=0.7, 
            color="blue"
        )
        ax.plot(
            df["year"], 
            df["minimum debt service coverage ratio"], 
            color="red", 
            marker="o", 
            label="Minimum DSCR Covenant", 
            linewidth=2
        )
        ax.set_xticks(df["year"].astype(int))
        ax.set_title("Debt Service Coverage Ratio (DSCR) Over Time", fontsize=14)
        ax.set_xlabel("Year", fontsize=12)
        ax.set_ylabel("Ratio", fontsize=12)
        ax.legend()
        ax.grid(visible=True, linestyle="--", alpha=0.5)
        return fig
    except ValueError as ve:
        st.error(str(ve))
        return None
    except Exception as e:
        st.error(f"Error creating chart: {e}")
        return None
    
# Function to handle memo generation
def generate_memo(is_material):
    try:
        memo_type = "Material Change Memo" if is_material else "Non-Material Change Memo"
        additional_content = ""
        chart_fig = None  # Initialize chart_fig

        # Parse uploaded files
        for uploaded_file in uploaded_files:
            if uploaded_file.type == "text/plain":
                additional_content += uploaded_file.read().decode("utf-8") + "\n"
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(uploaded_file)
                for para in doc.paragraphs:
                    additional_content += para.text + "\n"
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                df = pd.read_excel(uploaded_file)
                additional_content += df.to_string(index=False) + "\n"
                chart_fig = generate_dscr_chart(df)  # Generate chart
            elif uploaded_file.type == "text/csv":
                csv_df = pd.read_csv(uploaded_file)
                additional_content += csv_df.to_string(index=False) + "\n"
            else:
                additional_content += "Unsupported file type.\n"

        st.session_state["additional_content"] = additional_content
        st.session_state["chart_fig"] = chart_fig  # Store chart in session state

        # Prepare memo content
        document_content = f"Uploaded content:\n{additional_content}\nUser changes:\n{user_changes.strip()}"
        messages = [
            {
                "role": "user",
                "content": (
                    f"{document_content}\n\n---\n\n"
                    "Please act as a commercial lender at a top bank and so you must have very corporate communication. You closed a commercial loan in 2021, and after 3 years, the borrower has requested a 6-month extension to the loan term due to permitting issues that took longer than expected. To maintain a strong relationship with this borrower, you are incentivized to secure credit team approval for the loan term extension. Create a non-material change memo to formalize this 6-month extension. It is imperative that any output returned is in plain text form."
                    "The memo should have the following structure and include a visual chart to support the case. Ensure formatting is consistent, and headers for all sections are bold. It is important that you use bullet points for each new piece of informnation on each new line."
                    "Structure for the Memo:"
                    "Section 1: Background Information"
                    "Please return all information on separate lines. Include relevant property information, investment summary, and rationale for the initial investment, from the LP memo. Please list all of the information on separate lines."
                    "Section 2: Financial Information"
                    "Please include all relevant financial information that can be found in the uploaded PFS_CSV. This document is the personal financial statement. I want all information about the borrower to be included on separate lines."
                    "Section 3: Rationale for the Extension"
                    "Please ensure that the heading is on a separate line to the content. Please provide a summary of the uploaded email template pdf to describe the rationale behind the loan extension. Please ensure that there is no reference to inserting visual charts."
                    "When you sign off the memo, please ensure that it says: [name of the lender], [bank of the lender], [job title of the lender]"
                ),
            }
        ]
        response = client.chat.completions.create(model="gpt-4", messages=messages)
        st.session_state["generated_memo"] = response.choices[0].message.content

    except Exception as e:
        st.error(f"An error occurred: {e}")

# Buttons for memo generation and refresh
col1, col2, col3 = st.columns(3)
with col1:
    if st.button("Generate Non-Material Change Memo"):
        generate_memo(is_material=False)
with col2:
    if st.button("Generate Material Change Memo"):
        generate_memo(is_material=True)
with col3:
    if st.button("Refresh Page"):
        refresh_page()

# Display generated memo if available
if st.session_state.get("generated_memo", ""):
    st.subheader("Generated Memo")
    st.markdown(f"<div class='left-aligned'>{st.session_state['generated_memo']}</div>", unsafe_allow_html=True)

    # Save memo as a Word document
    output_doc = Document()
    output_doc.add_heading("Generated Memo", level=1)
    output_doc.add_paragraph(st.session_state["generated_memo"])

    buffer = BytesIO()
    output_doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Memo",
        data=buffer,
        file_name="Generated_Memo.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# Display chart after the memo
if st.session_state.get("chart_fig", None):
    st.pyplot(st.session_state["chart_fig"])